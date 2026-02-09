
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from io import BytesIO

def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', size=None, bold=False, color=None):
    """辅助函数：统一设置中西文字体、大小、粗细和颜色"""
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

def set_cell_bg(cell, color_hex):
    """设置单元格背景色 (Hex string, e.g., 'F2F2F2')"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_number(run):
    """添加页码域代码"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def generate_word_report(report):
    doc = Document()
    
    # 全局样式设置
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.25 # 1.25倍行距，阅读更舒适

    # 配置页眉页脚
    section = doc.sections[0]
    
    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "信息系统安全等级保护测评报告"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.style.font.size = Pt(9)
    # 给页眉加下边框
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 页脚 (页码)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run("第 ")
    set_font(run_f, font_name_cn='宋体', size=Pt(10))
    add_page_number(fp.add_run())
    run_f = fp.add_run(" 页")
    set_font(run_f, font_name_cn='宋体', size=Pt(10))

    # ==========================================
    # 封面页 (Cover Page)
    # ==========================================
    for _ in range(4): doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('信息系统安全等级保护测评报告')
    set_font(run, font_name_cn='黑体', size=Pt(26), bold=True)
    
    for _ in range(5): doc.add_paragraph()

    meta = report.get('metadata', {}) or {}
    timestamp_raw = report.get('timestamp', '')
    ts_code = timestamp_raw.replace('-', '').replace(' ', '').replace(':', '')[:14]
    report_code = f"DJBH-{report.get('timestamp', '')[:4]}-{ts_code}" if ts_code else "DJBH-202X-000000"
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    rows_data = [
        ('报告编号', report_code),
        ('系统/设备名称', meta.get('assetName', report.get('target'))),
        ('目标地址', report.get('target')),
        ('系统定级', meta.get('securityLevel', '三级')),
        ('物理位置', meta.get('location', '未知')),
        ('测评负责人', meta.get('evaluator', 'Admin')),
        ('测评时间', report.get('timestamp'))
    ]

    for idx, (label, val) in enumerate(rows_data):
        row = table.rows[idx]
        row.height = Cm(1.2) # 固定行高
        
        # 标签列
        cell0 = row.cells[0]
        cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell0.width = Cm(5)
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中对齐看起来更正式
        run0 = p0.add_run(label)
        set_font(run0, font_name_cn='宋体', bold=True)
        set_cell_bg(cell0, 'F9F9F9') # 浅灰背景

        # 内容列
        cell1 = row.cells[1]
        cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell1.width = Cm(10)
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 增加一点左缩进
        p1.paragraph_format.left_indent = Cm(0.5)
        run1 = p1.add_run(str(val))
        set_font(run1, font_name_cn='Times New Roman')

    doc.add_page_break()

    # ==========================================
    # 第一章：测评结果概述
    # ==========================================
    h1 = doc.add_heading(level=1)
    run = h1.add_run('第一章 测评结果概述')
    set_font(run, font_name_cn='黑体', size=Pt(16), bold=True, color=RGBColor(0,0,0))
    doc.add_paragraph()
    
    # 评分段落
    p_score = doc.add_paragraph()
    run = p_score.add_run("综合安全评分：")
    set_font(run, font_name_cn='宋体')
    run = p_score.add_run(f"{report.get('score', 0)} / 100")
    set_font(run, bold=True, size=Pt(14))
    
    score = report.get('score', 0)
    level = meta.get('securityLevel', '三级')
    compliance_txt = '基本符合' if score >= 80 else '不符合'
    risk_txt = '存在严重安全风险，必须立即进行整改加固。' if score < 60 else '仍存在部分安全隐患，建议按期整改。'
    
    p_desc = doc.add_paragraph()
    run = p_desc.add_run(f"系统安全防护能力{compliance_txt}等级保护{level}要求，{risk_txt}")
    set_font(run, font_name_cn='宋体')
    doc.add_paragraph()

    defects = report.get('defects', [])
    p_count = doc.add_paragraph()
    run = p_count.add_run(f"本次测评共发现 {len(defects)} 个安全问题：")
    set_font(run, font_name_cn='宋体', bold=True)

    # 风险统计表 (美化版)
    summary = report.get('summary', {})
    high, med, low = summary.get('high', 0), summary.get('medium', 0), summary.get('low', 0)
    total = len(defects) or 1 

    rtable = doc.add_table(rows=4, cols=4)
    rtable.style = 'Table Grid'
    rtable.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers = ['风险等级', '数量', '占比', '整改期限']
    for i, h in enumerate(headers):
        cell = rtable.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_bg(cell, '4F81BD') # 蓝色表头
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, font_name_cn='黑体', bold=True, color=RGBColor(255, 255, 255)) # 白字

    stats_data = [
        ('高危', high, '7天内', RGBColor(255, 0, 0)),
        ('中危', med, '30天内', RGBColor(255, 165, 0)),
        ('低危', low, '60天内', RGBColor(0, 0, 255))
    ]

    for idx, (label, count, limit, color) in enumerate(stats_data, 1):
        row = rtable.rows[idx]
        row.height = Cm(0.8)
        cells = row.cells
        
        # Helper to set content
        def set_cell_text(c, text, is_bold=False, text_color=None):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_font(r, font_name_cn='宋体', bold=is_bold, color=text_color)
        
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], count, True, color)
        set_cell_text(cells[2], f"{(count/total)*100:.1f}%")
        set_cell_text(cells[3], limit)

    doc.add_page_break()

    # ==========================================
    # 第二章：安全技术测评详情 (表格化布局)
    # ==========================================
    h1 = doc.add_heading(level=1)
    run = h1.add_run('第二章 安全技术测评详情')
    set_font(run, font_name_cn='黑体', size=Pt(16), bold=True, color=RGBColor(0,0,0))

    groups = {}
    for d in defects:
        parts = d.get('mlps_clause', '').split('-')
        domain = parts[1] if len(parts) >= 2 else '其他安全问题'
        if domain not in groups: groups[domain] = []
        groups[domain].append(d)

    domain_order = ['安全通信网络', '安全区域边界', '安全计算环境', '安全管理中心']
    sorted_domains = sorted(groups.keys(), key=lambda x: domain_order.index(x) if x in domain_order else 99)

    if not defects:
        doc.add_paragraph("本系统符合安全技术要求，未发现明显安全隐患。")

    for d_idx, domain in enumerate(sorted_domains, 1):
        doc.add_paragraph()
        h2 = doc.add_heading(level=2)
        run = h2.add_run(f"{d_idx}. {domain}")
        set_font(run, font_name_cn='黑体', size=Pt(14), bold=True, color=RGBColor(0,0,0))
        
        for i, d in enumerate(groups[domain], 1):
            doc.add_paragraph() # 两个漏洞之间的间距
            
            # 创建漏洞详情表 (6行2列)
            dt_table = doc.add_table(rows=6, cols=2)
            dt_table.style = 'Table Grid'
            dt_table.autofit = False 
            dt_table.allow_autofit = False
            
            # 设置列宽
            for r in dt_table.rows:
                r.cells[0].width = Cm(2.5)
                r.cells[1].width = Cm(13.5)

            # 行 0: 标题行 (合并单元格)
            row0 = dt_table.rows[0]
            cell_title = row0.cells[0]
            cell_title.merge(row0.cells[1])
            set_cell_bg(cell_title, 'F2F2F2') # 浅灰标题底色
            p = cell_title.paragraphs[0]
            run = p.add_run(f"问题 #{i}：{d.get('check_item')}")
            set_font(run, font_name_cn='黑体', bold=True, size=Pt(10.5))

            # 数据行定义
            rows_def = [
                ('风险等级', d.get('risk_level'), True),
                ('检测协议', d.get('protocol'), False),
                ('涉及条款', d.get('mlps_clause'), False),
                ('问题描述', d.get('description'), False),
                ('整改建议', d.get('suggestion'), False)
            ]

            for r_idx, (label, val, is_risk_row) in enumerate(rows_def, 1):
                row = dt_table.rows[r_idx]
                
                # 标签列
                c_label = row.cells[0]
                c_label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = c_label.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER # 标签居中
                run = p.add_run(label)
                set_font(run, font_name_cn='宋体', bold=True, size=Pt(10.5))
                
                # 内容列
                c_val = row.cells[1]
                c_val.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = c_val.paragraphs[0]
                run = p.add_run(str(val))
                
                color = None
                if is_risk_row:
                    if val == '高危': color = RGBColor(255, 0, 0)
                    elif val == '中危': color = RGBColor(255, 165, 0)
                    elif val == '低危': color = RGBColor(0, 0, 255)
                    set_font(run, font_name_cn='宋体', bold=True, color=color, size=Pt(10.5))
                else:
                    set_font(run, font_name_cn='宋体', size=Pt(10.5))

    doc.add_page_break()

    # ==========================================
    # 第三章：结论与整改
    # ==========================================
    h1 = doc.add_heading(level=1)
    run = h1.add_run('第三章 测评结论与整改要求')
    set_font(run, font_name_cn='黑体', size=Pt(16), bold=True, color=RGBColor(0,0,0))
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"经 NetAudit Pro v3.2 自动化测评，该信息系统的综合安全评分为 ")
    set_font(run, font_name_cn='宋体')
    run = p.add_run(f"{score} 分")
    set_font(run, font_name_cn='Times New Roman', bold=True)
    run = p.add_run("。")
    set_font(run, font_name_cn='宋体')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("整改要求：")
    set_font(run, font_name_cn='黑体', bold=True)
    
    reqs = [
        "高危问题必须在 7天内 完成整改（失陷系统需24小时内处理）",
        "中危问题必须在 30天内 完成整改",
        "低危问题必须在 60天内 完成整改",
        "整改完成后使用 NetAudit Pro 进行复测验证"
    ]
    
    for r_text in reqs:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(r_text)
        set_font(run, font_name_cn='宋体')

    doc.add_paragraph()
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run("--- 报告结束 ---")
    set_font(run, font_name_cn='宋体', color=RGBColor(128, 128, 128))

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f
